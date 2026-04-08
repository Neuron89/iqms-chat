"""
IQMS Chat — Web interface for querying IQMS ERP via Claude + MCP.
Users log in, ask questions in plain English, get answers from Claude
which queries the IQMS Oracle database through the MCP server.
"""

import csv
import io
import json
import logging
import os
import re
import subprocess
import tempfile
import uuid
from collections import deque
from datetime import datetime
from hashlib import sha256
from pathlib import Path
from functools import wraps

from flask import (
    Flask, render_template, request, redirect, url_for,
    session, flash, jsonify, send_from_directory,
)

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", os.urandom(32))
app.config["MAX_CONTENT_LENGTH"] = 20 * 1024 * 1024  # 20 MB

# ---------------------------------------------------------------------------
# In-memory log buffer (last 500 entries)
# ---------------------------------------------------------------------------
LOG_BUFFER: deque[dict] = deque(maxlen=500)


def _log(level: str, message: str, **extra):
    entry = {
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "level": level,
        "message": message,
        **extra,
    }
    LOG_BUFFER.append(entry)
    # Also print to stdout for systemd journal
    print(f"[{entry['timestamp']}] {level}: {message}", flush=True)


def log_info(message: str, **extra):
    _log("INFO", message, **extra)


def log_error(message: str, **extra):
    _log("ERROR", message, **extra)


def log_warn(message: str, **extra):
    _log("WARN", message, **extra)

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
AGENT_MEMORY_DIR = Path.home() / ".claude" / "agent-memory" / "erp-database"
IQMS_DOCS_DIR = Path.home() / "iqms-plugin" / "plugins" / "iqms-team-tools" / "data" / "iqms-docs"

# Key memory files loaded into every prompt (most important knowledge)
CORE_MEMORY_FILES = [
    "MEMORY.md",
    "schema-overview.md",
    "gotchas.md",
    "common-queries.md",
    "module-reference.md",
]

DATA_DIR = Path(__file__).parent / "data"
DATA_DIR.mkdir(exist_ok=True)
USERS_FILE = DATA_DIR / "users.json"
REPORTS_DIR = Path(__file__).parent / "reports"
REPORTS_DIR.mkdir(exist_ok=True)
UPLOADS_DIR = Path(__file__).parent / "uploads"
UPLOADS_DIR.mkdir(exist_ok=True)

MAX_UPLOAD_SIZE = 20 * 1024 * 1024  # 20 MB
ALLOWED_EXTENSIONS = {
    # Images (passed as file path for Claude to read visually)
    ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp",
    # Documents (text extracted and inlined)
    ".docx", ".xlsx", ".xls", ".csv",
    ".txt", ".md", ".json", ".xml", ".yaml", ".yml",
    ".pdf",
}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}

# MCP config — IQMS Oracle is always available, dataPARC only for Nycoa
IQMS_MCP = {
    "command": "node",
    "args": ["/home/hnester/iqms-plugin/plugins/iqms-team-tools/mcp-servers/iqms-oracle/index.js"],
    "env": {
        "IQMS_DB_USER": "mcp_readonly",
        "IQMS_DB_PASSWORD": "Palmvilla119",
        "IQMS_DB_CONNECT": "IQORA",
    },
}
DATAPARC_MCP = {
    "command": "node",
    "args": ["/home/hnester/dataparc-mcp/dist/index.js"],
    "env": {
        "SQL_SERVER": "10.10.1.248",
        "SQL_PORT": "1433",
        "SQL_DATABASE": "ctc_config",
        "SQL_USER": "hnester",
        "SQL_PASSWORD": "Palmvilla119",
        "SQL_DOMAIN": "scada",
        "SQL_TRUST_CERT": "true",
        "DEFAULT_PLANT": "NYCOA",
    },
}

# Write both config files (with and without dataPARC)
MCP_CONFIG_IQMS = DATA_DIR / "mcp_config_iqms.json"
MCP_CONFIG_IQMS.write_text(json.dumps({"mcpServers": {"iqms-oracle": IQMS_MCP}}, indent=2))

MCP_CONFIG_ALL = DATA_DIR / "mcp_config_all.json"
MCP_CONFIG_ALL.write_text(json.dumps({"mcpServers": {"iqms-oracle": IQMS_MCP, "dataparc": DATAPARC_MCP}}, indent=2))

CLAUDE_MODEL = os.environ.get("CLAUDE_MODEL", "opus")

# ePlant definitions
EPLANTS = {
    "1": {"name": "Shawsheen", "company": "Shawsheen Rubber Co Inc"},
    "2": {"name": "Nycoa", "company": "Nylon Corporation of America"},
    "3": {"name": "Bradford", "company": "Bradford Industries"},
}

def _load_core_memory() -> str:
    """Load key agent memory files for inclusion in the system prompt."""
    sections = []
    for fname in CORE_MEMORY_FILES:
        fpath = AGENT_MEMORY_DIR / fname
        if fpath.exists():
            try:
                content = fpath.read_text(encoding="utf-8", errors="replace")
                sections.append(f"=== {fname} ===\n{content}")
            except Exception:
                pass
    return "\n\n".join(sections)


SYSTEM_PROMPT_TEMPLATE = """You are an IQMS ERP database assistant. Users ask questions about manufacturing data,
production records, work orders, quality, inventory, and other ERP data stored in the IQMS Oracle database.

Use the iqms-oracle MCP tools to query the database and answer questions. Available IQMS tools:
- query: Run read-only SQL against the IQMS Oracle database
- list-tables: List tables matching a pattern
- describe-table: Show column details for a table
- sample-data: Preview rows from a table
- search-columns: Find columns by name across tables
- table-relationships: Show foreign key relationships
- table-indexes: Show indexes on a table
- row-count: Count rows in a table
{dataparc_section}
CRITICAL — ePlant Filtering:
The user is currently working with ePlant: {eplant_name} (EPLANT_ID = {eplant_id}, {eplant_company}).
You MUST add "WHERE EPLANT_ID = {eplant_id}" (or "AND EPLANT_ID = {eplant_id}") to EVERY query that
touches a table containing an EPLANT_ID column. Most major tables have this column (ARINVT, WORKORDER,
ORDERS, PDAYPROD, DAYPROD, ARCUSTO, etc. — 571+ tables). Always check if the table has EPLANT_ID and
filter accordingly. Never return data from other ePlants.

Important guidelines:
- All queries are read-only (SELECT only)
- The schema owner is IQMS — prefix tables as IQMS.TABLE_NAME
- PDAYPROD contains archived daily production data (the real historical data)
- DAYPROD is just the unarchived buffer (current/recent entries only)
- Give clear, concise answers. Format numbers and tables nicely using markdown.
- If you're unsure about a table or column, use the discovery tools first.
- Only show the final answer to the user. Do not describe your internal steps or tool calls.

Query Strategy — Break Down Complex Questions:
When a user asks a complex question (e.g. financial impact across all BOMs, cross-module analysis),
DO NOT try to answer it in one massive query. Instead:
1. First identify what specific data you need (which tables, which columns)
2. Run small, targeted queries one at a time — e.g. first find the item, then find BOMs using it,
   then calculate the impact
3. Use your loaded memory knowledge to pick the right tables immediately — don't waste time
   on discovery queries for tables you already know about
4. Keep individual SQL queries simple and fast — avoid huge JOINs across 5+ tables when you can
   run 2-3 focused queries and combine the results yourself
5. If a query returns too many rows, add ROWNUM or TOP limits and summarize

Report Generation:
When the user asks for a report, generate the data and format it as a clean, well-structured
markdown document. At the very end of your response, output the report content between these
exact markers on their own lines:

===REPORT_START===
(report filename, e.g. production_summary_2025.md)
===REPORT_CONTENT===
(full report content in markdown)
===REPORT_END===

Then tell the user their report is ready for download. The system will automatically detect
these markers and create a downloadable file. Use descriptive filenames with the ePlant name
and date. For CSV data, use .csv extension instead of .md.

KNOWLEDGE BASE & MEMORY:
You have access to an agent memory system at {agent_memory_dir}/ that contains institutional
knowledge about IQMS — schema details, gotchas, working query patterns, module references,
and past investigations. The key contents are loaded below.

You also have access to IQMS documentation at {iqms_docs_dir}/ — use the Read tool to consult
these docs when you need workflow/business logic not discoverable from the schema alone.

CRITICAL — Memory Updates:
After EVERY query where you discover something new, you MUST update the agent memory files.
This includes:
- New table/column relationships or meanings → update data-dictionary.md or schema-overview.md
- Query patterns that work well → update common-queries.md
- Gotchas, quirks, or traps → update gotchas.md
- New SQL patterns or functions → update sql-reference.md
- If appropriate, create a new .md file in the memory directory for a new topic.
Use the Edit tool to append to existing files. Keep entries concise. This is how you learn
and help future queries be faster and more accurate.

=== LOADED AGENT MEMORY ===
{agent_memory}
=== END AGENT MEMORY ===
"""

DATAPARC_PROMPT_SECTION = """
Additionally, you have access to dataPARC SCADA/historian tools for real-time and historical
process data. dataPARC is only used at the Nycoa plant. Available dataPARC tools:
- searchTags: Search for dataPARC tags by name pattern (e.g. %RX1%, %Temperature%)
- readCurrentValues: Get the current/latest value of one or more tags
- readRawTags: Read raw historical tag data for a time range
- readInterpolatedTags: Read interpolated (evenly-spaced) historical tag data
- readAtTimeTags: Read tag values at specific timestamps

When the user asks about process data, sensor readings, temperatures, pressures, line speeds,
or anything SCADA/historian related, use the dataPARC tools. You can combine IQMS and dataPARC
data in a single answer when relevant (e.g. correlating production records with process parameters).
"""

# Store conversation histories in memory (keyed by session chat_id)
conversations: dict[str, list] = {}

# ---------------------------------------------------------------------------
# User store
# ---------------------------------------------------------------------------

def _load_users() -> dict:
    if USERS_FILE.exists():
        return json.loads(USERS_FILE.read_text())
    return {}


def _save_users(users: dict):
    USERS_FILE.write_text(json.dumps(users, indent=2))


def _hash_pw(password: str, salt: str = "") -> str:
    return sha256(f"{salt}:{password}".encode()).hexdigest()


# ---------------------------------------------------------------------------
# Auth
# ---------------------------------------------------------------------------

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "username" not in session:
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return decorated


def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "username" not in session:
            return redirect(url_for("login"))
        users = _load_users()
        user = users.get(session["username"], {})
        if not user.get("is_admin"):
            flash("Admin access required.", "error")
            return redirect(url_for("chat"))
        return f(*args, **kwargs)
    return decorated


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username", "").strip().lower()
        password = request.form.get("password", "")

        users = _load_users()
        if username not in users:
            flash("Invalid credentials.", "error")
            return render_template("login.html")

        if users[username]["hash"] != _hash_pw(password, users[username]["salt"]):
            flash("Invalid credentials.", "error")
            return render_template("login.html")

        session["username"] = username
        session["display_name"] = users[username].get("display_name", username)
        session["is_admin"] = users[username].get("is_admin", False)
        session["chat_id"] = str(uuid.uuid4())
        session["eplant_id"] = "2"  # Default to Nycoa
        log_info(f"User '{username}' logged in")
        return redirect(url_for("chat"))

    return render_template("login.html")


@app.route("/logout")
def logout():
    chat_id = session.get("chat_id")
    if chat_id and chat_id in conversations:
        del conversations[chat_id]
    session.clear()
    return redirect(url_for("login"))


# ---------------------------------------------------------------------------
# ePlant selection
# ---------------------------------------------------------------------------

@app.route("/set-eplant", methods=["POST"])
@login_required
def set_eplant():
    data = request.get_json()
    eplant_id = data.get("eplant_id", "1")
    if eplant_id in EPLANTS:
        session["eplant_id"] = eplant_id
        return jsonify({"ok": True, "name": EPLANTS[eplant_id]["name"]})
    return jsonify({"error": "Invalid ePlant"}), 400


# ---------------------------------------------------------------------------
# File upload & processing
# ---------------------------------------------------------------------------

def _extract_text_from_docx(filepath: Path) -> str:
    """Extract text content from a .docx file."""
    from docx import Document
    doc = Document(str(filepath))
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_text_from_xlsx(filepath: Path) -> str:
    """Extract content from an .xlsx file as CSV-like text."""
    from openpyxl import load_workbook
    wb = load_workbook(str(filepath), read_only=True, data_only=True)
    output = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        output.append(f"=== Sheet: {sheet_name} ===")
        buf = io.StringIO()
        writer = csv.writer(buf)
        for row in ws.iter_rows(values_only=True):
            writer.writerow([str(c) if c is not None else "" for c in row])
        output.append(buf.getvalue())
    wb.close()
    return "\n".join(output)


def _process_upload(filepath: Path) -> dict:
    """Process an uploaded file. Returns dict with 'type' and content info."""
    ext = filepath.suffix.lower()

    if ext in IMAGE_EXTENSIONS:
        # Images: Claude CLI can read them via the Read tool
        return {"type": "image", "path": str(filepath), "name": filepath.name}

    if ext == ".pdf":
        # PDFs: Claude CLI can read them via the Read tool
        return {"type": "pdf", "path": str(filepath), "name": filepath.name}

    if ext == ".docx":
        text = _extract_text_from_docx(filepath)
        return {"type": "text", "content": text[:100000], "name": filepath.name}

    if ext in (".xlsx", ".xls"):
        text = _extract_text_from_xlsx(filepath)
        return {"type": "text", "content": text[:100000], "name": filepath.name}

    # Plain text / csv / json / etc — read directly
    try:
        text = filepath.read_text(encoding="utf-8", errors="replace")
        return {"type": "text", "content": text[:100000], "name": filepath.name}
    except Exception:
        return {"type": "error", "name": filepath.name}


@app.route("/upload", methods=["POST"])
@login_required
def upload():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "No file selected"}), 400

    ext = Path(f.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        return jsonify({"error": f"File type '{ext}' not supported"}), 400

    # Save to user-specific uploads directory
    username = session["username"]
    user_dir = UPLOADS_DIR / username
    user_dir.mkdir(exist_ok=True)

    # Unique filename to avoid collisions
    safe_name = re.sub(r'[^\w\-.]', '_', f.filename)
    unique_name = f"{uuid.uuid4().hex[:8]}_{safe_name}"
    filepath = user_dir / unique_name
    f.save(str(filepath))

    # Check file size
    if filepath.stat().st_size > MAX_UPLOAD_SIZE:
        filepath.unlink()
        return jsonify({"error": "File too large (max 20 MB)"}), 400

    # Process the file
    result = _process_upload(filepath)
    result["upload_id"] = unique_name
    result["original_name"] = f.filename

    log_info(f"File uploaded: {f.filename} ({ext})", user=username, type=result.get("type"))
    return jsonify(result)


# ---------------------------------------------------------------------------
# Chat
# ---------------------------------------------------------------------------

@app.route("/")
@login_required
def chat():
    chat_id = session.get("chat_id", "")
    messages = conversations.get(chat_id, [])
    eplant_id = session.get("eplant_id", "1")
    return render_template("chat.html",
                           username=session["username"],
                           display_name=session.get("display_name", session["username"]),
                           messages=messages,
                           eplants=EPLANTS,
                           current_eplant=eplant_id)


@app.route("/new-chat", methods=["POST"])
@login_required
def new_chat():
    old_id = session.get("chat_id")
    if old_id and old_id in conversations:
        del conversations[old_id]
    session["chat_id"] = str(uuid.uuid4())
    return redirect(url_for("chat"))


# Report marker pattern
REPORT_PATTERN = re.compile(
    r'===REPORT_START===\s*\n(.+?)\n===REPORT_CONTENT===\s*\n(.*?)\n===REPORT_END===',
    re.DOTALL,
)


def _extract_report(answer: str, username: str) -> tuple[str, str | None]:
    """Extract report from answer if present. Returns (cleaned_answer, report_url or None)."""
    match = REPORT_PATTERN.search(answer)
    if not match:
        return answer, None

    filename = match.group(1).strip()
    content = match.group(2).strip()

    # Sanitize filename
    filename = re.sub(r'[^\w\-.]', '_', filename)
    if not filename:
        filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"

    # Save to user-specific reports directory
    user_dir = REPORTS_DIR / username
    user_dir.mkdir(exist_ok=True)
    report_path = user_dir / filename
    report_path.write_text(content, encoding="utf-8")

    # Remove the report markers from the displayed answer
    clean_answer = answer[:match.start()].rstrip()
    report_url = url_for("download_report", filename=filename)

    return clean_answer, report_url


@app.route("/ask", methods=["POST"])
@login_required
def ask():
    data = request.get_json()
    question = data.get("question", "").strip()
    attachments = data.get("attachments", [])
    if not question:
        return jsonify({"error": "Empty question"}), 400

    chat_id = session.get("chat_id", str(uuid.uuid4()))
    if chat_id not in conversations:
        conversations[chat_id] = []

    username = session["username"]

    # Build display content for user message (show attachment names)
    display_content = question
    if attachments:
        file_names = [a.get("original_name", a.get("name", "file")) for a in attachments]
        display_content = question + "\n\n📎 " + ", ".join(file_names)

    # Add user message
    conversations[chat_id].append({
        "role": "user",
        "content": display_content,
        "timestamp": datetime.now().strftime("%H:%M"),
    })

    # Build the full prompt with file contents
    full_prompt = question

    # Collect image/pdf file paths for Claude to read
    file_paths_for_claude = []

    for att in attachments:
        att_type = att.get("type")
        name = att.get("original_name", att.get("name", "file"))

        if att_type == "text":
            # Inline text content
            content = att.get("content", "")
            full_prompt += f"\n\n--- Attached file: {name} ---\n{content}\n--- End of {name} ---"

        elif att_type in ("image", "pdf"):
            # Claude CLI will read these via the Read tool
            fpath = att.get("path", "")
            if fpath and Path(fpath).exists():
                file_paths_for_claude.append(fpath)
                full_prompt += f"\n\n[Attached {att_type}: {name} — saved at {fpath}. Use the Read tool to view it.]"

    # Build ePlant-specific system prompt
    eplant_id = session.get("eplant_id", "1")
    eplant = EPLANTS.get(eplant_id, EPLANTS["1"])

    # dataPARC only available for Nycoa (eplant 2)
    is_nycoa = eplant_id == "2"
    dataparc_section = DATAPARC_PROMPT_SECTION if is_nycoa else ""
    mcp_config_file = str(MCP_CONFIG_ALL) if is_nycoa else str(MCP_CONFIG_IQMS)

    # Load fresh agent memory for each query
    agent_memory = _load_core_memory()

    system_prompt = SYSTEM_PROMPT_TEMPLATE.format(
        eplant_id=eplant_id,
        eplant_name=eplant["name"],
        eplant_company=eplant["company"],
        dataparc_section=dataparc_section,
        agent_memory=agent_memory,
        agent_memory_dir=str(AGENT_MEMORY_DIR),
        iqms_docs_dir=str(IQMS_DOCS_DIR),
    )

    # Write system prompt to temp file (too large for CLI arg)
    prompt_file = DATA_DIR / f"prompt_{chat_id[:8]}.txt"
    prompt_file.write_text(system_prompt, encoding="utf-8")

    # Build the claude command
    cmd = [
        "claude",
        "-p",
        "--model", CLAUDE_MODEL,
        "--output-format", "json",
        "--mcp-config", mcp_config_file,
        "--permission-mode", "bypassPermissions",
        "--system-prompt-file", str(prompt_file),
        "--no-session-persistence",
        "--add-dir", str(AGENT_MEMORY_DIR),
        "--add-dir", str(IQMS_DOCS_DIR),
    ]

    # Add file paths as allowed directories so Claude can read them
    for fpath in file_paths_for_claude:
        cmd.extend(["--add-dir", str(Path(fpath).parent)])

    # Pass user's question via stdin (avoids CLI arg length limits)

    att_count = len(attachments)
    log_info(
        f"Query from '{username}': {question[:120]}{'...' if len(question) > 120 else ''}",
        user=username, eplant=eplant["name"], model=CLAUDE_MODEL,
        attachments=att_count, mcp="iqms+dataparc" if is_nycoa else "iqms",
    )

    start_time = datetime.now()
    try:
        result = subprocess.run(
            cmd,
            input=full_prompt,
            capture_output=True,
            text=True,
            timeout=300,
            cwd="/home/hnester",
        )
        elapsed = (datetime.now() - start_time).total_seconds()

        if result.returncode != 0:
            error_msg = result.stderr.strip() or "Claude returned an error."
            answer = f"Sorry, I encountered an error: {error_msg}"
            log_error(
                f"Claude CLI failed (exit {result.returncode}, {elapsed:.1f}s)",
                user=username, stderr=result.stderr[:500],
                stdout_preview=result.stdout[:500],
            )
        else:
            try:
                output = json.loads(result.stdout)
                answer = output.get("result", result.stdout.strip())
                # Log cost/usage info if available
                usage = {k: v for k, v in output.items() if k != "result"}
                log_info(
                    f"Query completed ({elapsed:.1f}s)",
                    user=username, **usage,
                )
            except json.JSONDecodeError:
                answer = result.stdout.strip()
                log_warn(
                    f"Non-JSON response ({elapsed:.1f}s)",
                    user=username, stdout_preview=result.stdout[:300],
                )

    except subprocess.TimeoutExpired:
        elapsed = (datetime.now() - start_time).total_seconds()
        answer = "Sorry, the query took too long. Try a more specific question."
        log_error(f"Query timed out after {elapsed:.1f}s", user=username)
    except Exception as e:
        elapsed = (datetime.now() - start_time).total_seconds()
        answer = f"Sorry, something went wrong: {str(e)}"
        log_error(f"Exception: {str(e)}", user=username)

    # Clean up temp prompt file
    try:
        prompt_file.unlink(missing_ok=True)
    except Exception:
        pass

    # Check for report in the answer
    answer, report_url = _extract_report(answer, username)
    if report_url:
        log_info(f"Report generated for '{username}'", report_url=report_url)

    # Add assistant message
    msg = {
        "role": "assistant",
        "content": answer,
        "timestamp": datetime.now().strftime("%H:%M"),
    }
    if report_url:
        msg["report_url"] = report_url
    conversations[chat_id].append(msg)

    resp = {"answer": answer}
    if report_url:
        resp["report_url"] = report_url
    return jsonify(resp)


# ---------------------------------------------------------------------------
# Logs (admin only)
# ---------------------------------------------------------------------------

@app.route("/logs")
@admin_required
def view_logs():
    return render_template("logs.html",
                           username=session["username"],
                           display_name=session.get("display_name", session["username"]))


@app.route("/api/logs")
@admin_required
def api_logs():
    level = request.args.get("level", "")
    user = request.args.get("user", "")
    limit = int(request.args.get("limit", 200))

    entries = list(LOG_BUFFER)
    if level:
        entries = [e for e in entries if e["level"] == level.upper()]
    if user:
        entries = [e for e in entries if e.get("user", "").lower() == user.lower()]
    entries = entries[-limit:]
    return jsonify(entries)


# ---------------------------------------------------------------------------
# Report downloads
# ---------------------------------------------------------------------------

@app.route("/reports/<filename>")
@login_required
def download_report(filename):
    username = session["username"]
    user_dir = REPORTS_DIR / username
    if not (user_dir / filename).exists():
        return "Report not found.", 404
    return send_from_directory(user_dir, filename, as_attachment=True)


# ---------------------------------------------------------------------------
# Export — convert markdown response to Excel / Word / PDF
# ---------------------------------------------------------------------------

def _md_to_html(md_text: str) -> str:
    """Convert markdown to styled HTML for PDF/Word rendering."""
    import markdown
    body = markdown.markdown(md_text, extensions=["tables", "fenced_code"])
    return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
  body {{ font-family: 'Segoe UI', Arial, sans-serif; font-size: 11pt; color: #1e293b;
         max-width: 800px; margin: 2rem auto; padding: 0 1rem; }}
  h1,h2,h3 {{ color: #92400e; border-bottom: 1px solid #e5e7eb; padding-bottom: 0.3rem; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1rem 0; font-size: 10pt; }}
  th {{ background: #f59e0b; color: #000; font-weight: 600; }}
  th, td {{ border: 1px solid #d1d5db; padding: 6px 10px; text-align: left; }}
  tr:nth-child(even) {{ background: #f9fafb; }}
  code {{ background: #f1f5f9; padding: 2px 5px; border-radius: 3px; font-size: 0.9em; }}
  pre {{ background: #f1f5f9; padding: 1rem; border-radius: 6px; overflow-x: auto; font-size: 0.85em; }}
  strong {{ color: #92400e; }}
</style></head><body>{body}</body></html>"""


def _parse_tables_from_md(md_text: str) -> list[list[list[str]]]:
    """Extract markdown tables as lists of rows. Each table is a list of rows, each row a list of cells."""
    tables = []
    current_table = []
    for line in md_text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            # Skip separator rows (|---|---|)
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if cells and not all(set(c) <= {"-", ":", " "} for c in cells):
                current_table.append(cells)
        else:
            if current_table:
                tables.append(current_table)
                current_table = []
    if current_table:
        tables.append(current_table)
    return tables


@app.route("/export", methods=["POST"])
@login_required
def export():
    data = request.get_json()
    content = data.get("content", "")
    fmt = data.get("format", "pdf")
    if not content:
        return jsonify({"error": "No content to export"}), 400

    username = session["username"]
    user_dir = REPORTS_DIR / username
    user_dir.mkdir(exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")

    try:
        if fmt == "pdf":
            import weasyprint
            html = _md_to_html(content)
            filename = f"export_{ts}.pdf"
            filepath = user_dir / filename
            weasyprint.HTML(string=html).write_pdf(str(filepath))

        elif fmt == "docx":
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.table import WD_TABLE_ALIGNMENT

            doc = Document()
            style = doc.styles["Normal"]
            style.font.name = "Calibri"
            style.font.size = Pt(11)

            tables = _parse_tables_from_md(content)
            table_idx = 0

            for line in content.split("\n"):
                stripped = line.strip()

                # Skip table rows (handled separately)
                if stripped.startswith("|") and stripped.endswith("|"):
                    cells = [c.strip() for c in stripped.strip("|").split("|")]
                    if all(set(c) <= {"-", ":", " "} for c in cells):
                        continue
                    # Check if this is the first row of a table we haven't rendered yet
                    if table_idx < len(tables) and cells == tables[table_idx][0]:
                        tbl_data = tables[table_idx]
                        table_idx += 1
                        tbl = doc.add_table(rows=len(tbl_data), cols=len(tbl_data[0]))
                        tbl.style = "Table Grid"
                        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                        for i, row_data in enumerate(tbl_data):
                            for j, cell_text in enumerate(row_data):
                                if j < len(tbl.columns):
                                    cell = tbl.rows[i].cells[j]
                                    cell.text = cell_text
                                    if i == 0:
                                        for run in cell.paragraphs[0].runs:
                                            run.bold = True
                        doc.add_paragraph()
                    continue

                if stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    doc.add_paragraph(stripped[2:], style="List Bullet")
                elif stripped:
                    # Handle bold markers
                    p = doc.add_paragraph()
                    parts = stripped.split("**")
                    for k, part in enumerate(parts):
                        if part:
                            run = p.add_run(part)
                            if k % 2 == 1:
                                run.bold = True

            filename = f"export_{ts}.docx"
            filepath = user_dir / filename
            doc.save(str(filepath))

        elif fmt == "xlsx":
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

            wb = Workbook()
            ws = wb.active
            ws.title = "Export"

            header_font = Font(name="Calibri", bold=True, color="000000", size=11)
            header_fill = PatternFill(start_color="F59E0B", end_color="F59E0B", fill_type="solid")
            thin_border = Border(
                left=Side(style="thin"), right=Side(style="thin"),
                top=Side(style="thin"), bottom=Side(style="thin"),
            )

            tables = _parse_tables_from_md(content)

            if tables:
                row_num = 1
                for t_idx, tbl in enumerate(tables):
                    if t_idx > 0:
                        row_num += 1  # blank row between tables

                    for i, row_data in enumerate(tbl):
                        for j, cell_text in enumerate(row_data):
                            cell = ws.cell(row=row_num, column=j + 1, value=cell_text)
                            cell.border = thin_border
                            cell.alignment = Alignment(wrap_text=True)
                            if i == 0:
                                cell.font = header_font
                                cell.fill = header_fill
                            # Try to convert numbers
                            try:
                                cell.value = float(cell_text.replace(",", "").replace("$", ""))
                                cell.number_format = '#,##0.00' if "." in cell_text else '#,##0'
                            except (ValueError, AttributeError):
                                pass
                        row_num += 1

                # Auto-fit column widths
                for col_cells in ws.columns:
                    max_len = max((len(str(c.value or "")) for c in col_cells), default=10)
                    ws.column_dimensions[col_cells[0].column_letter].width = min(max_len + 3, 50)
            else:
                # No tables found — dump text line by line
                for i, line in enumerate(content.split("\n"), 1):
                    ws.cell(row=i, column=1, value=line.strip())

            filename = f"export_{ts}.xlsx"
            filepath = user_dir / filename
            wb.save(str(filepath))

        else:
            return jsonify({"error": f"Unknown format: {fmt}"}), 400

        log_info(f"Export {fmt.upper()} for '{username}'", filename=filename)
        return jsonify({"url": url_for("download_report", filename=filename)})

    except Exception as e:
        log_error(f"Export failed: {str(e)}", user=username, format=fmt)
        return jsonify({"error": f"Export failed: {str(e)}"}), 500


# ---------------------------------------------------------------------------
# Admin — User Management
# ---------------------------------------------------------------------------

@app.route("/admin")
@admin_required
def admin():
    users = _load_users()
    return render_template("admin.html",
                           username=session["username"],
                           display_name=session.get("display_name", session["username"]),
                           users=users)


@app.route("/admin/add-user", methods=["POST"])
@admin_required
def add_user():
    username = request.form.get("username", "").strip().lower()
    display_name = request.form.get("display_name", "").strip()
    password = request.form.get("password", "")
    is_admin = request.form.get("is_admin") == "on"

    if not username or not password:
        flash("Username and password are required.", "error")
        return redirect(url_for("admin"))

    if len(password) < 4:
        flash("Password must be at least 4 characters.", "error")
        return redirect(url_for("admin"))

    users = _load_users()
    if username in users:
        flash(f"User '{username}' already exists.", "error")
        return redirect(url_for("admin"))

    salt = os.urandom(16).hex()
    users[username] = {
        "hash": _hash_pw(password, salt),
        "salt": salt,
        "display_name": display_name or username,
        "is_admin": is_admin,
        "created": datetime.now().isoformat(),
    }
    _save_users(users)
    flash(f"User '{display_name or username}' created.", "success")
    return redirect(url_for("admin"))


@app.route("/admin/delete-user", methods=["POST"])
@admin_required
def delete_user():
    username = request.form.get("username", "")
    if username == session["username"]:
        flash("You can't delete yourself.", "error")
        return redirect(url_for("admin"))

    users = _load_users()
    if username in users:
        name = users[username].get("display_name", username)
        del users[username]
        _save_users(users)
        flash(f"User '{name}' deleted.", "success")
    return redirect(url_for("admin"))


@app.route("/admin/reset-password", methods=["POST"])
@admin_required
def reset_password():
    username = request.form.get("username", "")
    new_password = request.form.get("new_password", "")

    if len(new_password) < 4:
        flash("Password must be at least 4 characters.", "error")
        return redirect(url_for("admin"))

    users = _load_users()
    if username not in users:
        flash("User not found.", "error")
        return redirect(url_for("admin"))

    salt = os.urandom(16).hex()
    users[username]["hash"] = _hash_pw(new_password, salt)
    users[username]["salt"] = salt
    _save_users(users)
    flash(f"Password reset for '{users[username].get('display_name', username)}'.", "success")
    return redirect(url_for("admin"))


# ---------------------------------------------------------------------------
# Bootstrap admin if no users exist
# ---------------------------------------------------------------------------
def _ensure_admin():
    users = _load_users()
    if not users:
        salt = os.urandom(16).hex()
        users["admin"] = {
            "hash": _hash_pw("admin", salt),
            "salt": salt,
            "display_name": "Administrator",
            "is_admin": True,
            "created": datetime.now().isoformat(),
        }
        _save_users(users)
        print(">>> Default admin account created (username: admin, password: admin)")
        print(">>> Change the password immediately via the admin panel!")


# ---------------------------------------------------------------------------
# Run
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    _ensure_admin()
    app.run(host="0.0.0.0", port=5055, debug=False)
